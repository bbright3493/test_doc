# -*- coding: utf-8 -*-
"""
__mktime__ = '2019/4/23 0023'
__author__ = 'Administrator'
__filename__ = 'docx'
文件注解：
"""

import os
import sys

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL, WD_ROW_HEIGHT_RULE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

global doc
global last_err
last_err = ""

#sys.path.append('..\\Lib\\site-packages')

def Init():
    global doc
    doc = Document()
    doc.styles["Normal"].font.name = u'宋体'
    doc.styles["Normal"].font.size = Pt(8)
    doc.styles["Normal"]._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    head = doc.add_heading("", level=1)
    run = head.add_run(u"社会主义核心价值观")
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(128, 0, 0)
    run.font.name = u"宋体"
    run._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    head.alignment = WD_ALIGN_PARAGRAPH.CENTER

def Save(dst):
    global doc
    global last_err
    try:
        doc.save(dst)
        return True
    except IOError as e:
        last_err = e.strerror
        return False

def AddLineTitle(title):
    global doc
    head = doc.add_heading("", level=2)
    run = head.add_run(title)
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0, 0, 0)
    run.font.name = u"宋体"
    run._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')

def AddInfoTitle(title):
    global doc
    head = doc.add_heading("", level=3)
    run = head.add_run(title)
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0, 0, 0)
    run.font.name = u"宋体"
    run._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')

def AddLineText(text):
    global doc
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(8)
    run.font.name = u"宋体"
    run._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')

def AddRecord(qj, cj, gh, mc, lx, ms, bz, tp):
    global doc
    table = doc.add_table(6, 2)
    table.style = "Table Grid"   # The First Table Style Of The MS Word

    for row in table.rows:
        row.height = Pt(14)
        row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
    table.rows[-1].height_rule = WD_ROW_HEIGHT_RULE.AUTO  # The Last Row Contains A Picture
    table.cell(0, 0).text = qj
    table.cell(0, 1).text = cj
    table.cell(0, 0).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    table.cell(0, 1).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    table.cell(1, 0).text = gh
    table.cell(1, 1).text = mc
    table.cell(1, 0).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    table.cell(1, 1).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    table.cell(2, 0).text = lx
    table.cell(2, 1).text = ms
    table.cell(2, 0).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    table.cell(2, 1).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    item = table.cell(3, 0).merge(table.cell(3, 1))
    item.text = bz
    item.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    item = table.cell(4, 0).merge(table.cell(4, 1))
    item.text = u"富强民主文明和谐爱国敬业诚信友善自由平等公正法治:"
    item.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    item = table.cell(5, 0).merge(table.cell(5, 1))
    p = item.paragraphs[0]
    p.add_run().add_picture(tp, Pt(300), Pt(300))
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

def GetLastError():
    global last_err
    return last_err


def Test():
    path = os.path.split(os.path.realpath(__file__))[0]
    file = path + "\\TEST.docx"

    Init()
    AddLineTitle(u"富强民主")
    AddInfoTitle(u"一、文明和谐")
    AddLineText(u"爱国敬业")
    AddRecord(u"诚信", u"友善", u"自由", u"平等", u"公正",
              u"法治", u"社会主义核心价值观", "2.jpg")
    Save(file)

Test()
print(GetLastError())
